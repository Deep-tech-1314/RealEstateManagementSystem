"""
Advanced script to directly update the Word document sections
"""
from docx import Document
from docx.shared import Pt
import re

doc_path = r'c:\Users\LENOVO\Downloads\Mini-Project Report-(Project-Title).docx'

try:
    doc = Document(doc_path)
    print("Document loaded successfully!")
except Exception as e:
    print(f"Error: {e}")
    exit(1)

# Project Title
PROJECT_TITLE = "Real Estate Management System"

# Update project title in headers/title
for paragraph in doc.paragraphs:
    text = paragraph.text.strip()
    if 'XXXX' in text or ('Project Title' in text and 'XXXX' in text):
        paragraph.text = PROJECT_TITLE

# Content sections
sections = {
    'Abstract': """The Real Estate Management System is a comprehensive web-based application developed using ASP.NET Core MVC framework and Entity Framework Core. The system provides an efficient platform for managing real estate properties, facilitating property listings, bookings, and inquiries. It enables property owners to list their properties with detailed information including images, location details, pricing, and amenities. Users can search and filter properties based on various criteria such as property type, location, price range, and listing type (sale/rent). The system incorporates role-based access control with separate interfaces for administrators and regular users. Administrators can manage all properties, users, bookings, and inquiries through a centralized dashboard. The system utilizes SQL Server as the database backend, ensuring data integrity and efficient data management. With features like property viewing history, booking management, and inquiry handling, the system streamlines the real estate business process, making it easier for both property owners and potential buyers/renters to interact and conduct transactions.""",
    
    'Introduction': """1.1 Background

Real estate management has traditionally been a complex process involving numerous intermediaries, paper-based documentation, and manual tracking of properties, clients, and transactions. The digital transformation of the real estate industry has become imperative to meet the growing demands of modern property buyers, sellers, and renters. With the increasing use of the internet and mobile devices, customers expect seamless access to property information, easy search capabilities, and efficient communication channels.

The need for an automated real estate management system arises from several challenges:
• Manual record keeping leads to errors and inefficiencies
• Difficulty in managing multiple properties across different locations
• Lack of centralized platform for property listings and inquiries
• Time-consuming processes for booking property viewings
• Challenges in tracking customer inquiries and responses

1.2 Problem Statement

Traditional real estate operations face significant challenges in managing property listings, handling customer inquiries, scheduling property viewings, and maintaining accurate records. The absence of a centralized digital platform results in:
• Inefficient property search and discovery process
• Poor communication between property owners and potential buyers/renters
• Lack of proper tracking mechanisms for bookings and inquiries
• Difficulty in managing property data and images
• Limited analytics and reporting capabilities

1.3 Objectives

The primary objectives of this Real Estate Management System project are:
• To develop a web-based platform for efficient property management
• To provide an intuitive interface for property search and filtering
• To enable seamless booking of property viewings
• To facilitate communication through inquiry management
• To implement role-based access control for administrators and users
• To ensure secure data management and user authentication
• To provide comprehensive administrative tools for property and user management
• To create a responsive and user-friendly interface

1.4 Scope

The system encompasses the following features:
• User registration and authentication system
• Property listing and management
• Advanced property search and filtering
• Property booking system for viewing appointments
• Inquiry management for customer queries
• Administrative dashboard with analytics
• User profile management
• Image upload and management for properties
• City-based property organization"""
}

# Function to find and update section
def update_section(section_name, content):
    found_section = False
    section_started = False
    
    for i, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        
        # Look for section heading
        if section_name.lower() in text.lower() and not found_section:
            found_section = True
            section_started = True
            continue
        
        # If section found and next paragraph is empty or ready for content
        if section_started:
            # If we hit another major heading, stop
            if text and any(heading in text for heading in ['Technology', 'Implementation', 'Conclusion', 'References']):
                break
            
            # Add content to empty paragraph or create new
            if not text or len(text) < 10:
                if not text:
                    paragraph.text = content
                    break
                else:
                    # Insert new paragraph after current
                    new_para = doc.paragraphs[i].insert_paragraph_before(content)
                    break

# Try to update sections
print("Attempting to update document sections...")

# Update Abstract
update_section('Abstract', sections['Abstract'])

# Update Introduction  
update_section('Introduction', sections['Introduction'])

# Save updated document
output_path = r'c:\Users\LENOVO\Downloads\Mini-Project Report-(Project-Title)-FILLED.docx'
try:
    doc.save(output_path)
    print(f"\nUpdated document saved to: {output_path}")
except Exception as e:
    print(f"\nNote: Could not auto-save document. Error: {e}")
    print("Please manually copy content from report_filled_content.txt")

print("\nFor best results, please:")
print("1. Open report_filled_content.txt")
print("2. Copy each section content")
print("3. Paste into corresponding sections in your Word document")

